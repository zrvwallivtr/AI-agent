import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch

import openpyxl
import docx

import src.tools.file_reader as file_reader_module
from src.tools.file_reader import (
    FileReader,
    _session_path,
    _read_file_in_dropbox
)
from src.agent.llm import LLM
from tools import file_reader


@pytest.fixture
def file_reader():
    return FileReader(session="test_session")

@pytest.fixture
def isolated_dropbox(tmp_path, monkeypatch):
    monkeypatch.setattr(file_reader_module, "DROPBOX_DIR", tmp_path)
    monkeypatch.setattr(file_reader_module, "MODEL", "mock_reader_model")
    monkeypatch.setattr(file_reader_module, "FILE_OR_NOT_PROMPT", "Mock File Prompt")
    monkeypatch.setattr(file_reader_module, "GET_FILE_LIST_PROMPT", "Mock List Prompt")

    return tmp_path


# ========================================
# Module function tests
# ========================================

def test_session_path_handling(isolated_dropbox):
    # Custom session directory
    session_path = _session_path("mock_session")
    assert session_path == isolated_dropbox / "mock_session"

    # Fallback default chat directory
    default_path = _session_path(None)
    assert default_path == isolated_dropbox / "chat"

def test_read_file_in_dropbox_helper(isolated_dropbox):
    # Create mock file in dropbox
    test_file = isolated_dropbox / "test_file.txt"
    test_file.write_text("Dropbox content", encoding="utf-8")
    
    assert _read_file_in_dropbox(test_file) == "Dropbox content"


# ========================================
# Parser method tests
# ========================================

def test_read_txt(file_reader, isolated_dropbox):
    # Create test file
    test_path = file_reader.store_file_path / "test.txt"
    file_reader.store_file_path.mkdir(parents=True, exist_ok=True)
    test_path.write_text("Plain text", encoding="utf-8")

    # Ensure file content is read
    assert file_reader._read_txt(test_path) == "Plain text"

def test_read_txt_error(file_reader, isolated_dropbox):
    # non-existing path
    missing_path = file_reader.store_file_path / "missing.txt"

    err_msg = file_reader._read_txt(missing_path)

    # Ensure error message is returned
    assert "Error reading txt file" in err_msg

def test_read_csv(file_reader, isolated_dropbox):
    # Create test file
    test_path = file_reader.store_file_path / "test.csv"
    file_reader.store_file_path.mkdir(parents=True, exist_ok=True)
    test_path.write_text("csv content", encoding="utf-8")

    assert file_reader._read_txt(test_path) == "csv content"

def test_read_xlsx(file_reader, isolated_dropbox, tmp_path):
    test_path = tmp_path / "test.xlsx"

    # Generate workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "TestSheet"
    ws.append(["Name", "Role", "Status"])
    ws.append(["Alice", "Developer", "Active"])
    ws.append([None, None, None]) # Empty row to test guardrail
    ws.append(["Bob", "Designer", None])
    wb.save(test_path)

    # Run method
    result = file_reader._read_xlsx(test_path)

    # Ensure correct output format
    assert "=== Sheet: TestSheet ===" in result
    assert "Name | Role | Status" in result
    assert "Alice | Developer | Active" in result
    assert "Bob | Designer | " in result

def test_read_non_existent_xlsx(file_reader, isolated_dropbox, tmp_path):
    non_existent_path = tmp_path / "non_existent.xlsx"
    result = file_reader._read_xlsx(non_existent_path)

    assert "Error reading xlsx file" in result

def test_read_yaml(file_reader, isolated_dropbox):
    # Create test file
    test_path = file_reader.store_file_path / "test.yaml"
    file_reader.store_file_path.mkdir(parents=True, exist_ok=True)
    test_path.write_text("yaml content", encoding="utf-8")

    assert file_reader._read_txt(test_path) == "yaml content"

def test_read_yml(file_reader, isolated_dropbox):
    # Create test file
    test_path = file_reader.store_file_path / "test.yml"
    file_reader.store_file_path.mkdir(parents=True, exist_ok=True)
    test_path.write_text("yml content", encoding="utf-8")

    assert file_reader._read_txt(test_path) == "yml content"

def test_read_toml(file_reader, isolated_dropbox):
    file_reader = FileReader(session="test_session")

    # Create test file
    test_path = file_reader.store_file_path / "test.toml"
    file_reader.store_file_path.mkdir(parents=True, exist_ok=True)
    test_path.write_text("toml content", encoding="utf-8")

    assert file_reader._read_txt(test_path) == "toml content"

def test_read_xml(file_reader, isolated_dropbox):
    # Create test file
    test_path = file_reader.store_file_path / "test.xml"
    file_reader.store_file_path.mkdir(parents=True, exist_ok=True)
    test_path.write_text("xml content", encoding="utf-8")

    assert file_reader._read_txt(test_path) == "xml content"

@patch("pdfplumber.open")
def test_read_pdf_extraction(mock_pdf_open, file_reader, isolated_dropbox):
    pdf_path    = file_reader.store_file_path / "test.pdf"

    # Mock extracted contents from non-existing pdf
    mock_page_1 = MagicMock()
    mock_page_1.extract_text.return_value = "Page One Content"
    mock_page_2 = MagicMock()
    mock_page_2.extract_text.return_value = "Page Two Content"

    mock_pdf_instance       = MagicMock()
    mock_pdf_instance.pages = [mock_page_1, mock_page_2]
    mock_pdf_open.return_value.__enter__.return_value = mock_pdf_instance

    # Ensure mock pdf content returned
    extracted_text = file_reader._read_pdf(pdf_path)
    assert extracted_text == "Page One Content\n\nPage Two Content"

def test_read_docx(file_reader, isolated_dropbox, tmp_path):
    test_path = tmp_path / "test.docx"

    # Generate simple docx file
    doc = docx.Document()
    doc.add_paragraph("First paragraph text block.")
    doc.add_paragraph("   ")  # Whitespace paragraph to test filtering
    doc.add_paragraph("Second paragraph text block.")
    doc.save(test_path)

    # Run method
    result = file_reader._read_docx(test_path)

    assert result == "First paragraph text block.\nSecond paragraph text block."

def test_read_docx_failure(file_reader, isolated_dropbox, tmp_path):
    non_existent_path = tmp_path / "corrupt.docx"

    # Create a garbage file that isn't a zip archive
    non_existent_path.write_text("Not a real zip/docx file structure")
    
    # Run method
    result = file_reader._read_docx(non_existent_path)

    assert "Error reading docx file" in result

@patch("src.tools.file_reader.epub.read_epub")
@patch("src.tools.file_reader.BeautifulSoup")
def test_read_epub(mock_bs_class, mock_read_epub, file_reader):
    # Mock the internal document items
    mock_item = MagicMock()
    mock_item.get_type.return_value = 9  # 9 corresponds to ebooklib.ITEM_DOCUMENT
    mock_item.get_content.return_value = b"<html><body>Chapter 1 content</body></html>"
    
    mock_book = mock_read_epub.return_value
    mock_book.get_items.return_value = [mock_item]
    
    # Mock the BeautifulSoup HTML extraction behavior
    mock_soup = mock_bs_class.return_value
    mock_soup.get_text.return_value = "Chapter 1 content"
    
    # Run method
    result = file_reader._read_epub(Path("dummy.epub"))
    
    # Verify the loop extracted data correctly
    assert result == "Chapter 1 content"

@patch("src.tools.file_reader.epub.read_epub")
def test_read_epub_failure(mock_read_epub, file_reader):
    mock_read_epub.side_effect = Exception("Invalid encryption keys or format")

    # Run method
    result = file_reader._read_epub(Path("broken.epub"))

    assert "Error reading EPUB file" in result

def test_read_code_blocks(file_reader, isolated_dropbox):
    # Create test file
    code_path = file_reader.store_file_path / "test_script.py"
    file_reader.store_file_path.mkdir(parents=True, exist_ok=True)
    code_path.write_text("print('hello')", encoding="utf-8")

    # Ensure reader returns correct language and format
    expected_output = "```py\nprint('hello')\n```"
    assert file_reader._read_code(code_path) == expected_output


# ========================================
# Storage tools
# ========================================

def test_store_content_in_dropbox(file_reader, isolated_dropbox):
    filename    = "cached_context.txt"
    
    result = file_reader._store_content_in_dropbox("Persisted Data Structures", filename)
    
    # Ensure file is stored in dropbox
    target_path = file_reader.store_file_path / filename
    assert target_path.exists()
    assert target_path.read_text(encoding="utf-8") == "Persisted Data Structures"
    assert result is None  # Successful implementation returns None

def test_list_available_files(file_reader, isolated_dropbox):
    file_reader.store_file_path = isolated_dropbox

    # Ensure directory doesn't exist yet
    assert file_reader.list_available_files() == []

    # Creates two different test file type
    file_reader.store_file_path.mkdir(parents=True, exist_ok=True)
    (file_reader.store_file_path / "test.txt").write_text("content")
    (file_reader.store_file_path / "test.py").write_text("content")
   
    # Ensure function lists all available files
    file_list = file_reader.list_available_files()
    assert len(file_list) == 2
    assert "test.txt" in file_list
    assert "test.py" in file_list

def test_clear_session_dropbox(file_reader, isolated_dropbox):
    file_reader.store_file_path.mkdir(parents=True, exist_ok=True)
   
    # Create test file
    file_path = file_reader.store_file_path / "test.txt"
    file_path.write_text("data")

    # Clear
    message = file_reader.clear_session_dropbox(session="clear_session")

    # Ensure removed file
    assert "Cleared all files" in message
    assert not file_path.exists()
    assert not file_reader.store_file_path.exists()


# ========================================
# LLM integration
# ========================================

def test_require_file(file_reader, isolated_dropbox, monkeypatch):
    # Mock model reponse
    mock_response = MagicMock()
    mock_llm_call = MagicMock(return_value=mock_response)
    monkeypatch.setattr(LLM, "response_with_new_sys_prompt_and_context", mock_llm_call)

    # Model returns 'True' without proper format
    # (test if the function is able to convert to proper format)
    mock_response.message.content = "  The answer is TRUE. Content needed. "
    assert file_reader.require_file_or_not([], "Prompt message") is True

def test_not_require_file(file_reader, isolated_dropbox, monkeypatch):
    # Mock model reponse
    mock_response = MagicMock()
    mock_llm_call = MagicMock(return_value=mock_response)
    monkeypatch.setattr(LLM, "response_with_new_sys_prompt_and_context", mock_llm_call)

    # Model returns 'False' with proper format
    mock_response.message.content = "False"
    assert file_reader.require_file_or_not([], "Prompt message") is False

def test_get_filenames_segregation(file_reader, isolated_dropbox, monkeypatch):
    file_reader.store_file_path.mkdir(parents=True, exist_ok=True)

    # Generate an active file
    (file_reader.store_file_path / "present.txt").write_text("here")

    # Simulate erratic formatting with markdown bullets, commas, quotes, and newlines,
    # only one file exist in the mock model response list.
    mock_response                   = MagicMock()
    mock_response.message.content   = "* present.txt\n- 'absent.csv', \"missing.py\""
    
    mock_llm_call = MagicMock(return_value=mock_response)
    monkeypatch.setattr(LLM, "response_with_new_sys_prompt_and_context", mock_llm_call)

    found, not_found = file_reader.get_filenames([], "This message is irrelevant in this test.")

    # Ensure only existing file is returned
    assert "present.txt" in found
    assert "absent.csv" in not_found
    assert "missing.py" in not_found


# ========================================
# Execution
# ========================================

def test_read_files(file_reader, isolated_dropbox, monkeypatch):
    file_reader.store_file_path.mkdir(parents=True, exist_ok=True)

    # Create files
    (file_reader.store_file_path / "test.txt").write_text("Text file content", encoding="utf-8")
    (file_reader.store_file_path / "test.py").write_text("Python code content", encoding="utf-8")

    # Mock user message
    mock_user_message = {"role": "user", "content": "This message is irrelevant in this test."}
    monkeypatch.setattr(LLM, "user", MagicMock(return_value=mock_user_message))

    # Mock LLM message
    mock_llm_response = MagicMock(return_value=("Mock model read code contents and response to message.", 150, 45))
    monkeypatch.setattr(file_reader_module.LLM, "model_response", mock_llm_response)

    # Execute function
    context_in      = [{"role": "system", "content": "System prompt"}]
    files_to_read   = ["test.txt", "test.py"]
    response, prompt_tokens, out_tokens = file_reader.read_files(context_in, files_to_read, "Read relevant files.")

    # Ensure proper returned values
    assert response         == "Mock model read code contents and response to message."
    assert prompt_tokens    == 150
    assert out_tokens       == 45

    # History checker for LLM.user() in 'read_files'
    args, _         = file_reader_module.LLM.user.call_args
    bundled_prompt  = args[0]
    
    # Ensure all related file are included into the message
    assert "Filename = test.txt" in bundled_prompt
    assert "Text file content" in bundled_prompt
    assert "Filename = test.py" in bundled_prompt
    assert "```py\nPython code content\n```" in bundled_prompt  # Ensured routed via code block parser
    assert "User input: Read relevant files." in bundled_prompt
